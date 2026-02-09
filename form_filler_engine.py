"""
Vendor Form Auto-Filler — Core Engine
=======================================
Reads ANY vendor registration form (xlsx, docx, pdf), understands its layout,
maps fields to your master company data (.md), and fills it in the original format.

SETUP:
    pip install anthropic openpyxl python-docx pdfplumber pypdf reportlab

USAGE:
    from form_filler_engine import fill_form
    fill_form("master_data.md", "blank_form.xlsx", "filled_form.xlsx")
"""

import json
import os
import re
import subprocess
import sys
import tempfile
from pathlib import Path

import anthropic

# ── Claude client (uses ANTHROPIC_API_KEY env var) ────────────────────────────
client = anthropic.Anthropic()
MODEL = "claude-sonnet-4-20250514"


# =============================================================================
#  STEP 1: Read the form — extract text/structure regardless of format
# =============================================================================

def read_form(file_path: str) -> dict:
    """
    Reads a form file and returns its text content + metadata.
    Supports: .xlsx, .docx, .pdf
    """
    ext = Path(file_path).suffix.lower()

    if ext == ".xlsx":
        return _read_xlsx(file_path)
    elif ext in (".docx", ".doc"):
        return _read_docx(file_path)
    elif ext == ".pdf":
        return _read_pdf(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Supported: xlsx, docx, pdf")


def _read_xlsx(path):
    from openpyxl import load_workbook
    wb = load_workbook(path)
    sheets = {}
    for name in wb.sheetnames:
        ws = wb[name]
        rows = []
        for row in ws.iter_rows(min_row=1, max_row=ws.max_row, max_col=ws.max_column, values_only=False):
            row_data = {}
            for cell in row:
                if cell.value is not None:
                    row_data[cell.coordinate] = str(cell.value)
            if row_data:
                rows.append(row_data)
        sheets[name] = rows
    return {"format": "xlsx", "sheets": sheets, "raw_text": _sheets_to_text(sheets)}


def _sheets_to_text(sheets):
    lines = []
    for name, rows in sheets.items():
        lines.append(f"--- Sheet: {name} ---")
        for row in rows:
            lines.append("  ".join(f"{k}: {v}" for k, v in row.items()))
    return "\n".join(lines)


def _read_docx(path):
    import docx
    doc = docx.Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract from tables
    table_data = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_data.append(cells)

    raw_text = "\n".join(paragraphs)
    if table_data:
        raw_text += "\n\n--- Tables ---\n"
        for row in table_data:
            raw_text += " | ".join(row) + "\n"

    return {"format": "docx", "paragraphs": paragraphs, "tables": table_data, "raw_text": raw_text}


def _read_pdf(path):
    import pdfplumber
    pages_text = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            pages_text.append(text)

            # Also try tables
            tables = page.extract_tables()
            if tables:
                for table in tables:
                    for row in table:
                        if row:
                            pages_text.append(" | ".join(str(c or "") for c in row))

    return {"format": "pdf", "raw_text": "\n".join(pages_text)}


# =============================================================================
#  STEP 2: Use Claude to map form fields → master data
# =============================================================================

MAPPING_PROMPT = """You are a precise data extraction and mapping assistant.

I have a MASTER DATA FILE with my company's information (the source of truth):
---
{master_data}
---

I have a BLANK VENDOR REGISTRATION FORM with these fields/labels:
---
{form_text}
---

Your job:
1. Identify every field/label in the form that needs to be filled in.
2. Match each form field to the correct value from the master data.
3. Return a JSON array of mappings.

For XLSX forms, return:
```json
[
  {{"form_cell": "B3", "form_label": "Company Name", "value": "Weston Manus", "confidence": "high"}},
  ...
]
```

For DOCX forms, return:
```json
[
  {{"form_label": "Company Name", "placeholder": "___________", "value": "Weston Manus", "confidence": "high"}},
  ...
]
```

For PDF forms, return:
```json
[
  {{"form_label": "Company Name", "value": "Weston Manus", "confidence": "high"}},
  ...
]
```

Rules:
- Return ONLY valid JSON, no markdown, no explanation.
- Match intelligently: "Vendor Name" = "Company Name", "Tax ID" = "VAT / GST Number", etc.
- If a form asks for something not in the master data, set value to null and confidence to "none".
- If a form field is a header/title/instruction (not a fillable field), skip it.
- For Yes/No fields, match to the master data value.
- The "confidence" field should be "high", "medium", or "none".
- Be thorough — don't miss any fillable field.
- For the format field, include "format": "{form_format}" in your response.
"""


def map_fields(master_data: str, form_content: dict) -> list:
    """Use Claude to intelligently map form fields to master data values."""
    response = client.messages.create(
        model=MODEL,
        max_tokens=4000,
        messages=[{
            "role": "user",
            "content": MAPPING_PROMPT.format(
                master_data=master_data,
                form_text=form_content["raw_text"],
                form_format=form_content["format"]
            ),
        }],
    )

    raw = response.content[0].text.strip()
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1].rsplit("```", 1)[0]
    return json.loads(raw)


# =============================================================================
#  STEP 3: Fill the form in its original format
# =============================================================================

def fill_xlsx(input_path: str, output_path: str, mappings: list):
    """Fill an Excel form using cell references from the mapping."""
    from openpyxl import load_workbook
    from openpyxl.utils import coordinate_to_tuple
    wb = load_workbook(input_path)
    ws = wb.active

    # Build set of cells we need to fill
    cells_to_fill = {}
    for m in mappings:
        cell = m.get("form_cell")
        value = m.get("value")
        if cell and value and value != "null":
            cells_to_fill[cell] = value

    # Unmerge any merged ranges that contain cells we need to fill
    for mc in list(ws.merged_cells.ranges):
        for cell_ref in cells_to_fill:
            row, col = coordinate_to_tuple(cell_ref)
            if (mc.min_row <= row <= mc.max_row and mc.min_col <= col <= mc.max_col
                    and not (row == mc.min_row and col == mc.min_col)):
                # This cell is merged and NOT the top-left anchor — unmerge the range
                ws.unmerge_cells(str(mc))
                break

    # Now fill
    filled = 0
    for cell, value in cells_to_fill.items():
        ws[cell] = value
        filled += 1

    wb.save(output_path)
    return filled


def fill_docx(input_path: str, output_path: str, mappings: list):
    """Fill a Word form by replacing placeholders or blank lines next to labels."""
    import docx
    from docx.shared import Pt
    doc = docx.Document(input_path)

    # Build a lookup: lowercase label → value
    # Also build alternate keys for fuzzy matching
    label_map = {}
    for m in mappings:
        label = m.get("form_label", "").strip().lower().rstrip(":")
        value = m.get("value")
        if label and value and value != "null":
            label_map[label] = value
            # Also add without common suffixes/prefixes
            clean = re.sub(r'\s*\(.*?\)\s*', '', label).strip()
            if clean and clean != label:
                label_map[clean] = value

    print(f"   📝 Label map has {len(label_map)} entries")

    filled = 0

    # Strategy 1: Fill table cells (most common Word form layout)
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            cells = row.cells
            # Build list of unique cells (skip merged duplicates)
            unique_cells = []
            seen_tcs = []
            for cell in cells:
                tc = cell._tc
                if tc not in seen_tcs:
                    seen_tcs.append(tc)
                    unique_cells.append(cell)

            for i, cell in enumerate(unique_cells):
                cell_text = cell.text.strip().lower().rstrip(":")
                if not cell_text:
                    continue

                # Also try matching without parenthetical notes
                cell_text_clean = re.sub(r'\s*\(.*?\)\s*', '', cell_text).strip().rstrip(":")

                matched_value = None
                # Direct match
                if cell_text in label_map:
                    matched_value = label_map[cell_text]
                elif cell_text_clean and cell_text_clean in label_map:
                    matched_value = label_map[cell_text_clean]
                else:
                    # Fuzzy: check if any label is contained in cell text or vice versa
                    for lbl, val in label_map.items():
                        if lbl in cell_text or cell_text in lbl:
                            matched_value = val
                            break

                if matched_value and i + 1 < len(unique_cells):
                    answer_cell = unique_cells[i + 1]
                    # Skip if answer cell already has meaningful content that looks like a label
                    existing = answer_cell.text.strip()
                    if existing and ':' in existing:
                        continue

                    # Clear the answer cell
                    for paragraph in answer_cell.paragraphs:
                        for run in paragraph.runs:
                            run.text = ""

                    # Write the value
                    first_para = answer_cell.paragraphs[0]
                    if first_para.runs:
                        first_para.runs[0].text = matched_value
                    else:
                        run = first_para.add_run(matched_value)
                        run.font.size = Pt(9)

                    filled += 1
                    print(f"   ✅ T{t_idx}R{r_idx}: '{cell_text[:30]}' → '{matched_value[:30]}'")

    # Strategy 2: Fill inline "Label: ____" patterns in paragraphs
    for paragraph in doc.paragraphs:
        for label, value in label_map.items():
            pattern = re.compile(
                re.escape(label) + r'\s*:\s*[_\s]*$',
                re.IGNORECASE
            )
            if pattern.search(paragraph.text):
                new_text = re.sub(
                    r'(:\s*)[_\s]*$',
                    f'\\1{value}',
                    paragraph.text,
                    flags=re.IGNORECASE
                )
                if paragraph.runs:
                    for run in paragraph.runs:
                        run.text = ""
                    paragraph.runs[0].text = new_text
                else:
                    paragraph.add_run(new_text)
                filled += 1

    doc.save(output_path)
    print(f"   📝 Docx saved with {filled} fills")
    return filled


def fill_pdf(input_path: str, output_path: str, mappings: list):
    """
    Fill a PDF form. Tries fillable fields first, then falls back to text overlay.
    """
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(input_path)

    # Check if the PDF has fillable form fields
    if reader.get_fields():
        return _fill_pdf_fields(input_path, output_path, mappings, reader)
    else:
        return _fill_pdf_overlay(input_path, output_path, mappings)


def _fill_pdf_fields(input_path, output_path, mappings, reader):
    """Fill a PDF that has actual fillable form fields."""
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.append(reader)

    label_map = {}
    for m in mappings:
        label = m.get("form_label", "").strip().lower()
        value = m.get("value")
        if label and value and value != "null":
            label_map[label] = value

    fields = reader.get_fields()
    filled = 0
    for field_name, field_obj in fields.items():
        field_lower = field_name.strip().lower()
        for label, value in label_map.items():
            if label in field_lower or field_lower in label:
                writer.update_page_form_field_values(writer.pages[0], {field_name: value})
                filled += 1
                break

    with open(output_path, "wb") as f:
        writer.write(f)
    return filled


def _fill_pdf_overlay(input_path, output_path, mappings):
    """
    For non-fillable PDFs: detect table cell boundaries with pdfplumber,
    then overlay text into the correct cells using reportlab.
    """
    import pdfplumber
    from reportlab.pdfgen import canvas as rl_canvas
    from pypdf import PdfReader, PdfWriter
    import io

    # Build label → value map
    label_map = {}
    for m in mappings:
        label = m.get("form_label", "").strip().lower().rstrip(":")
        value = m.get("value")
        if label and value and value != "null":
            label_map[label] = value
            clean = re.sub(r'\s*\(.*?\)\s*', '', label).strip().rstrip(":")
            if clean and clean != label:
                label_map[clean] = value

    # Find field positions using table cell boundaries
    field_positions = []

    with pdfplumber.open(input_path) as pdf:
        for page_idx, page in enumerate(pdf.pages):
            page_height = page.height
            page_width = page.width
            tables = page.find_tables()
            words = page.extract_words(keep_blank_chars=False, extra_attrs=["size"])

            for table in tables:
                for row in table.rows:
                    cells = row.cells
                    if not cells:
                        continue

                    # Get text content for each cell
                    cell_texts = []
                    for cell_bbox in cells:
                        if cell_bbox is None:
                            cell_texts.append(("", None))
                            continue
                        cx0, ctop, cx1, cbottom = cell_bbox
                        cell_words = [w for w in words
                                      if w["x0"] >= cx0 - 2 and w["x1"] <= cx1 + 2
                                      and w["top"] >= ctop - 2 and w["bottom"] <= cbottom + 2]
                        text = " ".join(w["text"] for w in sorted(cell_words, key=lambda w: (w["top"], w["x0"])))
                        cell_texts.append((text.strip(), cell_bbox))

                    # Match labels to next empty cell
                    for i, (text, bbox_i) in enumerate(cell_texts):
                        if not text or not bbox_i:
                            continue
                        text_lower = text.lower().rstrip(":").strip()
                        text_clean = re.sub(r'\s*\(.*?\)\s*', '', text_lower).strip()

                        for label, value in label_map.items():
                            matched = False
                            if text_lower == label or text_clean == label:
                                matched = True
                            elif label in text_lower and len(label) > 3:
                                matched = True

                            if matched:
                                for j in range(i + 1, len(cell_texts)):
                                    next_text, next_bbox = cell_texts[j]
                                    if next_bbox is None:
                                        continue
                                    if len(next_text) < 3 or next_text == text:
                                        cx0, ctop, cx1, cbottom = next_bbox
                                        fill_x = cx0 + 4
                                        fill_y = ctop + 3

                                        # Check for duplicate positions
                                        dupe = any(abs(fp["y"] - fill_y) < 5 and abs(fp["x"] - fill_x) < 5
                                                   for fp in field_positions)
                                        if not dupe:
                                            field_positions.append({
                                                "page": page_idx,
                                                "x": fill_x,
                                                "y": fill_y,
                                                "value": value,
                                                "font_size": 8,
                                                "page_height": page_height,
                                                "page_width": page_width,
                                            })
                                        break
                                break

    if not field_positions:
        # Fallback to guide
        print("   ⚠️  Could not detect field positions, generating fill guide")
        guide_path = output_path.replace(".pdf", "_fill_guide.txt")
        lines_out = ["VENDOR FORM FILL GUIDE", "=" * 50, ""]
        filled = 0
        for m in mappings:
            lbl = m.get("form_label", "Unknown")
            val = m.get("value", "")
            if val and val != "null":
                lines_out.append(f"  {lbl}: {val}")
                filled += 1
            else:
                lines_out.append(f"  {lbl}: [NOT FOUND]")
        with open(guide_path, "w") as f:
            f.write("\n".join(lines_out))
        import shutil
        shutil.copy2(input_path, output_path)
        return filled

    # Create overlay and merge
    reader = PdfReader(input_path)
    writer = PdfWriter()

    pages_fields = {}
    for fp in field_positions:
        pages_fields.setdefault(fp["page"], []).append(fp)

    for page_idx, page in enumerate(reader.pages):
        if page_idx in pages_fields:
            packet = io.BytesIO()
            ph = pages_fields[page_idx][0]["page_height"]
            pw = pages_fields[page_idx][0]["page_width"]
            c = rl_canvas.Canvas(packet, pagesize=(pw, ph))

            for fp in pages_fields[page_idx]:
                rl_y = ph - fp["y"] - fp["font_size"] - 2
                c.setFont("Helvetica", fp["font_size"])
                c.drawString(fp["x"], rl_y, fp["value"])

            c.save()
            packet.seek(0)
            overlay = PdfReader(packet)
            page.merge_page(overlay.pages[0])

        writer.add_page(page)

    with open(output_path, "wb") as f:
        writer.write(f)

    filled = len(field_positions)
    print(f"   📝 PDF overlay: placed {filled} text fields")
    return filled


# =============================================================================
#  MAIN: Orchestrator
# =============================================================================

def fill_form(master_data_path: str, form_path: str, output_path: str) -> dict:
    """
    Main entry point. Reads master data, reads the form, maps fields, fills it.

    Args:
        master_data_path: Path to your master .md file
        form_path: Path to the blank vendor form (xlsx, docx, or pdf)
        output_path: Where to save the filled form

    Returns:
        dict with results summary
    """
    # Load master data
    master_data = Path(master_data_path).read_text(encoding="utf-8")

    # Step 1: Read the form
    print(f"📄 Reading form: {form_path}")
    form_content = read_form(form_path)
    fmt = form_content["format"]
    print(f"   Format: {fmt}")

    # Step 2: Map fields using Claude
    print(f"🤖 Mapping form fields to master data...")
    mappings = map_fields(master_data, form_content)
    total_fields = len(mappings)
    fillable = sum(1 for m in mappings if m.get("value") and m["value"] != "null")
    print(f"   Found {total_fields} fields, {fillable} can be filled from master data")

    # Step 3: Fill the form
    print(f"✍️  Filling form...")
    if fmt == "xlsx":
        filled = fill_xlsx(form_path, output_path, mappings)
    elif fmt == "docx":
        filled = fill_docx(form_path, output_path, mappings)
    elif fmt == "pdf":
        filled = fill_pdf(form_path, output_path, mappings)

    print(f"✅ Done! Filled {filled} fields → {output_path}")

    # Log any fields that couldn't be filled
    missing = [m for m in mappings if not m.get("value") or m["value"] == "null"]
    if missing:
        print(f"\n⚠️  {len(missing)} fields not found in master data:")
        for m in missing:
            print(f"   - {m.get('form_label', 'Unknown field')}")

    return {
        "output_path": output_path,
        "format": fmt,
        "total_fields": total_fields,
        "filled": filled,
        "missing": [m.get("form_label") for m in missing],
        "mappings": mappings,
    }


# ── CLI ───────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Fill any vendor form from master data")
    parser.add_argument("--master", required=True, help="Path to master .md file")
    parser.add_argument("--form", required=True, help="Path to blank vendor form")
    parser.add_argument("--output", help="Output path (default: auto-generated)")
    args = parser.parse_args()

    if not args.output:
        stem = Path(args.form).stem
        ext = Path(args.form).suffix
        args.output = f"{stem}_FILLED{ext}"

    fill_form(args.master, args.form, args.output)
